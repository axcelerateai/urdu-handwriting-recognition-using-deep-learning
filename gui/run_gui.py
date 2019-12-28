import sys
import time
import tempfile
import os
import glob
import shutil
import zipfile
import gui.uhw_rc
import pyinsane2
import gui.abtform as abtform
import gui.termsform as termform
import gui.docsform as docsform
from gui.uiform import Ui_MainWindow
from PySide2.QtUiTools import QUiLoader
from PySide2 import QtCore, QtGui, QtWidgets, QtPrintSupport
from PySide2.QtCore import QFile, QTextStream, Qt, Signal, Slot
from PySide2.QtWidgets import QFileDialog, QMessageBox, QPushButton
from docx import Document

from gui.waitingspinnerwidget import QtWaitingSpinner
from pdf2image import convert_from_path, convert_from_bytes

from run_model import create_and_run_model

cnn_rnn_ctc_greedy_search       = "gui/configs/CNN_RNN_CTC_GS.json"
cnn_rnn_ctc_beam_search         = "gui/configs/CNN_RNN_CTC_BS.json"
cnn_rnn_ctc_beam_search_with_lm = "gui/configs/CNN_RNN_CTC_BLS.json"
encoder_decoder_greedy_search   = "gui/configs/Encoder_Decoder_GS.json"
encoder_decoder_beam_search     = "gui/configs/Encoder_Decoder_BS.json"

# Need to set icon tick below too. 
default_config_file = cnn_rnn_ctc_beam_search_with_lm

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.process_thread = ProcessingThread()
        self.process_thread.completed_signal.connect(self.thread_finished)
        self.process_thread.halted_signal.connect(self.halted_ocr)
        self.process_thread.model_loaded_signal.connect(self.loading_complete)
        # Load model in thread
        self.is_queued = bool(False)
        self.is_model_loaded = bool(False)
        self.is_config_changed = bool(False)
        self.process_thread.start()

        QtGui.QFontDatabase.addApplicationFont(":/fonts/jnoori.ttf")
        QtGui.QFontDatabase.addApplicationFont(":/fonts/cousine.ttf")

        self.statusBar().showMessage("Select image(s)")
        self.ui.textBrowser.setReadOnly(True)
        self.ui.textBrowser.setAlignment(Qt.AlignRight)

        # Spinner setup
        self.spinner = QtWaitingSpinner(self.ui.spinner_widget)

        self.light_theme()

        self.current_pdf = ""
        self.current_archive = ""
        self.currentpath = ""
        self.list_of_images = []
        self.currpic = 0
        self.totalpics = 0
        self.tempdir = tempfile.mkdtemp()
        self.result_pages_list_g = []
        self.isSplit = bool(False)
        self.selected_folder = "./"
        self.selected_save_path = "./"

        print(self.tempdir)

        self.tick_icon = self.style().standardIcon(
            getattr(QtWidgets.QStyle, "SP_DialogApplyButton")
        )
        self.noIcon = QtGui.QIcon()
        self.createMenu()

        self.ui.prev_button.hide()
        self.ui.next_button.hide()

        self.hide_textBrowser()

        self.ui.prev_button.setIcon(QtGui.QIcon(":/icon/left-arrow.png"))
        self.ui.next_button.setIcon(QtGui.QIcon(":/icon/right-arrow.png"))

        # connects
        self.ui.pushButton.clicked.connect(self.on_pushButton_clicked)
        self.ui.proc_button.clicked.connect(self.on_proc_button_clicked)
        self.ui.save_button.clicked.connect(self.on_save_button_clicked)
        self.ui.next_button.clicked.connect(self.nextImage)
        self.ui.prev_button.clicked.connect(self.backImage)
        self.ui.save_word_button.clicked.connect(self.save_as_word)
        self.ui.select_folder_button.clicked.connect(self.open_multiple_direc)
        self.ui.textBrowser.textChanged.connect(self.handleTextChanged)
        self.ui.save_pdf_button.clicked.connect(self.save_as_pdf)
        self.ui.stop_proc_button.clicked.connect(self.stop_ocr)
        self.ui.split_pages_button.clicked.connect(self.split_by_pages)
        self.ui.open_archive_button.clicked.connect(self.open_archive)

        # Font for output
        font = QtGui.QFont("Jameel Noori Nastaleeq", 17, 1)
        self.ui.textBrowser.setFont(font)

        self.ui.label_3.installEventFilter(self)
        self.ui.imgScrollArea.installEventFilter(self)

        self.handleTextChanged()

        # Thread for Scanning
        self.scan_thread = ScanningThread()
        self.scan_thread.tempdir = self.tempdir
        self.scan_thread.scan_done.connect(self.scanning_complete)

        # Make ribbon
        self.ui.toolBar.addWidget(self.ui.pushButton)
        self.ui.pushButton.setToolTip("Select Image(s)")
        self.ui.toolBar.addWidget(self.ui.select_folder_button)
        self.ui.select_folder_button.setToolTip("Select Folder(s)")
        self.ui.toolBar.addWidget(self.ui.open_archive_button)
        self.ui.open_archive_button.setToolTip("Select Archive (.zip) file")
        self.ui.toolBar.addSeparator()
        self.start_action = self.ui.toolBar.addWidget(self.ui.proc_button)
        self.run_action = self.ui.proc_button.setToolTip("Start OCR")
        self.stop_action = self.ui.toolBar.addWidget(self.ui.stop_proc_button)
        self.ui.stop_proc_button.setToolTip("Stop OCR")
        self.ui.toolBar.addSeparator()
        self.ui.toolBar.addWidget(self.ui.save_button)
        self.ui.save_button.setToolTip("Save as txt file")
        self.ui.toolBar.addWidget(self.ui.save_word_button)
        self.ui.save_word_button.setToolTip("Save as docx file")
        self.ui.toolBar.addWidget(self.ui.save_pdf_button)
        self.ui.save_pdf_button.setToolTip("Save as pdf file")
        self.ui.toolBar.addSeparator()
        self.ui.toolBar.addWidget(self.ui.split_pages_button)
        self.ui.split_pages_button.setToolTip("Split text to individual pages")

        self.stop_action.setVisible(False)

        # New scrollable view for selected images
        self.ui.imgScrollArea.setWidgetResizable(True)
        self.scrollable_imagelabels = ImagesScrollView(self)
        self.single_image_hide()
        self.resize

    def closeEvent(self, event):
        print("Exiting")
        self.cleandirectory(self.tempdir)
        os.removedirs(self.tempdir)

    def cleandirectory(self, path):
        folder = path
        for the_file in os.listdir(folder):
            file_path = os.path.join(folder, the_file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(e)

    def setIcons(self, color):
        # color = "white"
        self.changeicon(self.ui.save_button, "&#61686", color)
        self.changeicon(self.ui.save_word_button, "&#61890", color)
        self.changeicon(self.ui.proc_button, "&#61573", color)
        self.changeicon(self.ui.pushButton, "&#61893", color)
        self.changeicon(self.ui.select_folder_button, "&#58878", color)
        self.changeicon(self.ui.save_pdf_button, "&#61889", color)
        self.changeicon(self.ui.stop_proc_button, "&#62093", color)
        self.changeicon(self.ui.split_pages_button, "&#62970", color)
        self.changeicon(self.ui.open_archive_button, "&#62480", color)

    def changeicon(self, button, icon, color):
        icontext = QtGui.QTextDocument()
        icontext.setHtml(
            '<p align=center><font color="' + color + '">' + icon + ";</p>"
        )
        icontext.setDefaultFont(QtGui.QFont("Cousine Nerd Font Mono", 18))
        pixmap = QtGui.QPixmap(icontext.size().width(), icontext.size().height())
        pixmap.fill(Qt.transparent)
        painter = QtGui.QPainter(pixmap)
        icontext.drawContents(painter)
        buttonicon = QtGui.QIcon(pixmap)
        button.setIcon(buttonicon)
        button.setIconSize(pixmap.rect().size())
        painter.end()

    def disable(self, qobj):
        qobj.setEnabled(False)

    def enable(self, qobj):
        qobj.setEnabled(True)

    def trial(self, fileName):
        if fileName[0]:
            self.selected_folder = os.path.dirname(fileName[0][0])

        if fileName[0] and fileName[1] == "Image Files (*.png *.jpg *.jpeg *.bmp)":
            self.newImageList(fileName[0])

        elif fileName[0] and fileName[1] == "PDF files (*.pdf )":
            self.open_pdf(fileName[0])

    # Open/Select images
    def on_pushButton_clicked(self):
        fileName = QFileDialog.getOpenFileNames(
            self,
            ("Open Image"),
            self.selected_folder,
            ("Image Files (*.png *.jpg *.jpeg *.bmp);;PDF files (*.pdf )"),
        )

        if fileName[0]:
            self.selected_folder = os.path.dirname(fileName[0][0])

        if fileName[0] and fileName[1] == "Image Files (*.png *.jpg *.jpeg *.bmp)":
            self.newImageList(fileName[0])

        elif fileName[0] and fileName[1] == "PDF files (*.pdf )":
            self.open_pdf(fileName[0])

    def newImageList(self, list):
        if len(list) == 0:
            QtWidgets.QMessageBox.critical(
                self,
                "No images selected",
                "You have not selected any images.",
                QtWidgets.QMessageBox.Ok,
            )
            return
        self.single_image_hide()
        if len(list) == 1:
            self.disable(self.ui.next_button)
        else:
            self.enable(self.ui.next_button)
        self.currpic = 0
        self.list_of_images = list
        if not (self.currentpath == list[0] and len(list) == 1):
            self.ui.textBrowser.setPlainText("")
            self.hide_textBrowser()
            self.currentpath = list[0]
        self.totalpics = len(list) - 1

        self.disp_scrollable_imglist(list)
        self.displayImage(list[self.currpic])
        self.displayImageStatus()
        self.statusBar().showMessage("Ready")

    def disp_scrollable_imglist(self, list):
        self.scrollable_imagelabels.dispImages(
            list, self.ui.imgScrollArea.width(), self.ui.imgScrollArea.height()
        )
        self.ui.imgScrollArea.setWidget(self.scrollable_imagelabels)

    def open_pdf(self, fileName):
        self.cleandirectory(self.tempdir)
        if not (self.current_pdf == fileName[0] and len(fileName) == 1):
            self.ui.textBrowser.setPlainText("")
            self.hide_textBrowser()
            self.current_pdf = fileName[0]

        new_list = []
        i = 1
        for pdfname in fileName:
            imagespd = convert_from_path(pdfname)
            save_dir = self.tempdir
            for page in imagespd:
                new_name = os.path.join(save_dir, str(i) + ".png")
                page.save(new_name, "PNG")
                new_list.append(new_name)
                i += 1
        self.newImageList(new_list)

    def on_proc_button_clicked(self):
        if not self.list_of_images:
            QtWidgets.QMessageBox.critical(
                self,
                "No Images Selected",
                "You have not selected any images.",
                QtWidgets.QMessageBox.Ok,
            )
            return

        self.spinner.start()
        self.statusBar().showMessage("Processing")
        self.hide_textBrowser()

        if not self.is_queued:
            self.process_thread.image_list = self.list_of_images

        self.start_action.setVisible(False)
        self.stop_action.setVisible(True)

        if self.is_model_loaded:
            self.process_thread.start()
        else:
            self.is_queued = bool(True)

    def thread_finished(self, result_list):
        self.is_queued = bool(False)
        self.spinner.stop()
        self.start_action.setVisible(True)
        self.stop_action.setVisible(False)

        self.isSplit = bool(False)
        self.result_pages_list_g = result_list
        result_combined = "".join(result_list)
        self.statusBar().showMessage("Done")
        self.show_textBrowser()
        self.ui.textBrowser.setPlainText(result_combined)
        self.sync_split()

    def split_by_pages(self):
        if not self.isSplit:
            self.isSplit = bool(True)
            self.ui.textBrowser.setPlainText(self.result_pages_list_g[self.currpic])
            self.changeicon(self.ui.split_pages_button, "&#64087", "black")
            self.ui.split_pages_button.setToolTip("Combine text")
            self.single_image_show()

        else:
            self.isSplit = bool(False)
            result_combined = "".join(self.result_pages_list_g)
            self.ui.textBrowser.setPlainText(result_combined)
            self.changeicon(self.ui.split_pages_button, "&#62970", "black")
            self.ui.split_pages_button.setToolTip("Split text to individual pages")

            self.single_image_hide()

    def single_image_show(self):
        self.ui.label_3.show()
        self.ui.prev_button.show()
        self.ui.next_button.show()
        self.ui.curr_pic_label.show()
        self.ui.imgScrollArea.hide()

    def single_image_hide(self):
        self.ui.label_3.hide()
        self.ui.prev_button.hide()
        self.ui.next_button.hide()
        self.ui.curr_pic_label.hide()
        self.ui.imgScrollArea.show()

    def hide_textBrowser(self):
        self.ui.textBrowser.hide()
        self.ui.spinner_widget.show()
        self.ui.split_pages_button.setEnabled(False)

    def show_textBrowser(self):
        self.ui.textBrowser.show()
        self.ui.spinner_widget.hide()
        if len(self.list_of_images) > 1:
            self.ui.split_pages_button.setEnabled(True)

    def on_save_button_clicked(self):
        savepath = QFileDialog.getSaveFileName(
            self,
            ("Save "),
            os.path.join(self.selected_save_path, "new.txt"),
            ("Text (*.txt)"),
        )
        if savepath[0]:
            self.selected_save_path = savepath[0]
            savefile = savepath[0]  # + ".txt"
            outputext = self.ui.textBrowser.toPlainText()
            with open(savefile, "w+") as text_file:
                print(f"{outputext}", file=text_file)

    def save_as_word(self):
        savepath = QFileDialog.getSaveFileName(
            self,
            ("Save "),
            os.path.join(self.selected_save_path, "new.docx"),
            ("Word files (*.docx)"),
        )
        if savepath[0]:
            self.selected_save_path = savepath[0]

            outputext = self.ui.textBrowser.toPlainText()
            word_document = Document()
            p = word_document.add_paragraph(outputext)
            savefile = savepath[0]  # + ".docx"
            word_document.save(savefile)

    def eventFilter(self, widget, event):
        if event.type() == QtCore.QEvent.Resize:  
            selectedimg = QtGui.QPixmap(self.currentpath)
            self.disp_scrollable_imglist(self.list_of_images)
            self.ui.label_3.setPixmap(
                selectedimg.scaledToWidth(self.ui.imgScrollArea.width())
            )
            self.ui.label_3.setAlignment(Qt.AlignTop)
            return True
        return QtWidgets.QMainWindow.eventFilter(self, widget, event)

    def displayImage(self, imagePath):
        selectedimg = QtGui.QPixmap(imagePath)
        self.ui.label_3.setPixmap(
            selectedimg.scaledToWidth(self.ui.imgScrollArea.width())
        )
        self.ui.label_3.setAlignment(Qt.AlignTop)
        if self.currpic == 0:
            self.disable(self.ui.prev_button)

    def displayImageStatus(self):
        status = "Image {0} of {1}".format(self.currpic + 1, self.totalpics + 1)
        self.ui.curr_pic_label.setText(status)

    def nextImage(self):
        if self.currpic < self.totalpics:
            self.currpic = self.currpic + 1
            self.displayImage(self.list_of_images[self.currpic])
            self.displayImageStatus()
        if self.currpic == self.totalpics:
            self.disable(self.ui.next_button)
        if self.currpic > 0:
            self.enable(self.ui.prev_button)
        if self.isSplit:
            self.ui.textBrowser.setPlainText(self.result_pages_list_g[self.currpic])

    def backImage(self):
        if self.currpic > 0:
            self.currpic = self.currpic - 1
            self.displayImage(self.list_of_images[self.currpic])
            self.displayImageStatus()
        if self.currpic == 0:
            self.disable(self.ui.prev_button)
        if self.currpic < self.totalpics:
            self.enable(self.ui.next_button)
        if self.isSplit:
            self.ui.textBrowser.setPlainText(self.result_pages_list_g[self.currpic])

    def createMenu(self):
        # Print action
        self.printAction = QtWidgets.QAction("&Print", self)
        self.printAction.setShortcut("Ctrl+P")
        self.printAction.setStatusTip("Print document")
        self.printAction.triggered.connect(self.printCall)

        # Preview action
        self.previewAction = QtWidgets.QAction("&Print Preview", self)
        self.previewAction.setShortcut("Ctrl+F2")
        self.previewAction.setStatusTip("See what the print result will look like")
        self.previewAction.triggered.connect(self.previewCall)

        # Scan action
        self.scanAction = QtWidgets.QAction("&Scan", self)
        self.scanAction.setStatusTip("Scan a page")
        self.scanAction.triggered.connect(self.scan_page)

        # Exit action
        self.exitAction = QtWidgets.QAction("&Exit", self)
        self.exitAction.setShortcut("Ctrl+Q")
        self.exitAction.setStatusTip("Exit Application")
        self.exitAction.triggered.connect(self.exit)

        # Terms of use action
        self.termsAction = QtWidgets.QAction("&Terms and Conditions", self)
        self.termsAction.setStatusTip("View Terms and Conditions")
        self.termsAction.triggered.connect(self.terms)

        # Documentation action
        self.docsAction = QtWidgets.QAction("&How to Use?", self)
        self.docsAction.setStatusTip("How to Use?")
        self.docsAction.triggered.connect(self.docs)
        self.docsAction.setShortcut("Ctrl+H")

        # About action
        self.aboutAction = QtWidgets.QAction("&About", self)
        self.aboutAction.setStatusTip("About Urdu OCR")

        self.aboutAction.triggered.connect(self.about)
        # Open from single directory
        self.singdAction = QtWidgets.QAction("&Images", self)
        self.singdAction.setStatusTip("Select file(s)")
        self.singdAction.triggered.connect(self.on_pushButton_clicked)
        self.singdAction.setShortcut("Ctrl+I")

        # Open from multiple directories + subdirectories
        self.submultdAction = QtWidgets.QAction("&Folders", self)
        self.submultdAction.setStatusTip("Select folder(s)")
        self.submultdAction.triggered.connect(self.open_submultiple_direc)
        self.submultdAction.setShortcut("Ctrl+F")

        # Open Archive
        self.oarchiveAction = QtWidgets.QAction("&Archive (ZIP)", self)
        self.oarchiveAction.setStatusTip("Select archive (ZIP file)")
        self.oarchiveAction.triggered.connect(self.open_archive)

        # Light theme
        self.lightAction = QtWidgets.QAction("&Light Theme", self)
        self.lightAction.setStatusTip("Set a light theme")
        self.lightAction.triggered.connect(self.light_theme)

        # Dark theme
        self.darkAction = QtWidgets.QAction("&Dark Theme", self)
        self.darkAction.setStatusTip("Set a dark theme")
        self.darkAction.triggered.connect(self.dark_theme)

        # Option 11
        self.opt11 = QtWidgets.QAction("&Greedy Search", self)
        # self.opt11.setIcon(QtGui.QIcon())
        self.opt11.setStatusTip("Decoder type: Greedy Search")
        self.opt11.triggered.connect(self.method_opt11)

        # Option 12
        self.opt12 = QtWidgets.QAction("&Beam Search", self)
        self.opt12.setStatusTip("Decoder type: Beam Search")
        self.opt12.triggered.connect(self.method_opt12)

        # Option 13
        self.opt13 = QtWidgets.QAction("&Beam Search with Language Modeling", self)
        self.opt13.setStatusTip("Decoder type: Beam Search with Language Modeling")
        # Set this as default option
        self.opt13.setIcon(self.tick_icon)
        self.opt13.triggered.connect(self.method_opt13)

        # Option 21
        self.opt21 = QtWidgets.QAction("&Greedy Search", self)
        self.opt21.setStatusTip("Decoder type: Greedy Search")
        self.opt21.triggered.connect(self.method_opt21)

        # Option 22
        self.opt22 = QtWidgets.QAction("&Beam Search", self)
        self.opt22.setStatusTip("Decoder type: Beam Search")
        self.opt22.triggered.connect(self.method_opt22)

        # Custom config
        self.cust_conf = QtWidgets.QAction("&Custom Config", self)
        self.cust_conf.setStatusTip("Select a custom configuration file (.json format)")
        self.cust_conf.triggered.connect(self.change_config)

        self.model_opts = (
            self.opt11,
            self.opt12,
            self.opt13,
            self.opt21,
            self.opt22,
            self.cust_conf,
        )
        # Menu bar
        menuBar = self.menuBar()
        fileMenu = menuBar.addMenu("&File")
        openMenu = menuBar.addMenu("&Open")
        prefMenu = menuBar.addMenu("&Configuration")
        helpMenu = menuBar.addMenu("&Help")

        # File menu
        fileMenu.addAction(self.printAction)
        fileMenu.addAction(self.previewAction)
        fileMenu.addAction(self.scanAction)
        fileMenu.addSeparator()
        fileMenu.addAction(self.exitAction)
        # Open Menu
        openMenu.addAction(self.singdAction)
        openMenu.addAction(self.submultdAction)
        openMenu.addAction(self.oarchiveAction)

        conf_submenu1 = prefMenu.addMenu("&CNN_RNN_CTC")
        conf_submenu1.setStatusTip("Use CNN_RNN_CTC architecture with default settings")
        conf_submenu1.addAction(self.opt11)
        conf_submenu1.addAction(self.opt12)
        conf_submenu1.addAction(self.opt13)

        conf_submenu2 = prefMenu.addMenu("&Encoder-Decoder")
        conf_submenu2.setStatusTip("Use Encoder-Decoder architecture with default settings")
        conf_submenu2.addAction(self.opt21)
        conf_submenu2.addAction(self.opt22)
        prefMenu.addSeparator()
        prefMenu.addAction(self.cust_conf)
        # Help Menu
        helpMenu.addAction(self.docsAction)
        helpMenu.addSeparator()
        helpMenu.addAction(self.termsAction)
        helpMenu.addAction(self.aboutAction)

    def light_theme(self):
        file = QFile(":/stylesheet/light.txt")
        file.open(QFile.ReadOnly | QFile.Text)
        stream = QTextStream(file)
        self.setStyleSheet(stream.readAll())
        self.setIcons("black")
        self.spinner.setColor(QtGui.QColor(26, 25, 28))

    def dark_theme(self):
        file = QFile(":/stylesheet/dark.txt")
        file.open(QFile.ReadOnly | QFile.Text)
        stream = QTextStream(file)
        self.setStyleSheet(stream.readAll())
        self.setIcons("white")
        self.spinner.setColor(QtGui.QColor(206, 206, 206))

    def printCall(self):
        dialog = QtPrintSupport.QPrintDialog()
        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            self.ui.textBrowser.document().print_(dialog.printer())

    def previewCall(self):
        dialog = QtPrintSupport.QPrintPreviewDialog()
        dialog.paintRequested.connect(self.ui.textBrowser.print_)
        dialog.exec_()

    def scan_page(self):
        self.ui.textBrowser.setPlainText("")
        self.hide_textBrowser()

        self.cleandirectory(self.tempdir)
        temp_img_name = "1.png"
        temp_img_path = os.path.join(self.tempdir, temp_img_name)
        new_list = []
        self.scanAction.setEnabled(False)
        self.statusBar().showMessage("Scanning")
        pyinsane2.init()
        self.scan_thread.start()

    # Executed when textbrowser changes content
    def handleTextChanged(self):
        isenable = not self.ui.textBrowser.document().isEmpty()
        self.printAction.setEnabled(isenable)
        self.previewAction.setEnabled(isenable)
        if isenable:
            self.enable(self.ui.save_button)
            self.enable(self.ui.save_word_button)
            self.enable(self.ui.save_pdf_button)
        else:
            self.disable(self.ui.save_button)
            self.disable(self.ui.save_word_button)
            self.disable(self.ui.save_pdf_button)

        self.ui.textBrowser.setCursor(QtGui.QCursor(QtCore.Qt.IBeamCursor))

    def exit(self):
        sys.exit()

    def terms(self):
        t = TermsWindow()
        t.exec_()

    def docs(self):
        d = DocsWindow()
        d.exec_()

    def about(self):
        about = AboutWindow()
        about.exec_()

    def open_multiple_direc(self):
        input_dir = QFileDialog.getExistingDirectory(
            None, "Select a folder:", self.selected_folder
        )
        if input_dir:
            self.selected_folder = os.path.dirname(os.path.normpath(input_dir))
            self.ui.textBrowser.setPlainText("")
            self.hide_textBrowser()
            self.currentpath = input_dir
            files = []
            for ext in ("*.bmp", "*.png", "*.jpg", "*.jpeg"):
                files.extend(glob.glob(os.path.join(input_dir, ext)))
            self.newImageList(files)

    def open_submultiple_direc(self):
        input_dir = QFileDialog.getExistingDirectory(
            None, "Select a folder:", self.selected_folder
        )
        if input_dir:
            self.selected_folder = os.path.dirname(os.path.normpath(input_dir))
            self.ui.textBrowser.setPlainText("")
            self.hide_textBrowser()
            files = []
            for ext in ("**/*.bmp", "**/*.png", "**/*.jpg", "**/*.jpeg"):
                files.extend(glob.glob(os.path.join(input_dir, ext), recursive=True))
            self.newImageList(files)

    def open_archive(self):
        self.cleandirectory(self.tempdir)
        fileName = QFileDialog.getOpenFileName(
            self, ("Open Archive"), self.selected_folder, ("Zip Files (*.zip)")
        )[0]
        if fileName:
            self.selected_folder = os.path.dirname(fileName)
            with zipfile.ZipFile(fileName, "r") as zip_ref:
                zip_ref.extractall(self.tempdir)
            files = []
            for ext in ("*.bmp", "*.png", "*.jpg", "*.jpeg"):
                files.extend(glob.glob(os.path.join(self.tempdir, ext)))
            self.newImageList(files)

    def save_as_pdf(self):
        filename = QFileDialog.getSaveFileName(
            self,
            ("Save "),
            os.path.join(self.selected_save_path, "./new.pdf"),
            ("PDF (*.pdf)"),
        )[0]
        if filename:
            self.selected_save_path = filename
            printer = QtPrintSupport.QPrinter(QtPrintSupport.QPrinter.HighResolution)
            printer.setPageSize(QtPrintSupport.QPrinter.A4)
            printer.setColorMode(QtPrintSupport.QPrinter.Color)
            printer.setOutputFormat(QtPrintSupport.QPrinter.PdfFormat)
            printer.setOutputFileName(filename)
            self.ui.textBrowser.document().print_(printer)

    def stop_ocr(self):
        self.statusBar().showMessage("Stopping ... Please wait")
        self.spinner.setRevolutionsPerSecond(0.0001)
        self.process_thread.halt()

    def halted_ocr(self):
        self.spinner.stop()
        self.start_action.setVisible(True)
        self.stop_action.setVisible(False)

        self.isSplit = bool(False)
        self.statusBar().showMessage("Stopped")
        self.spinner.setRevolutionsPerSecond(1.57)

    def loading_complete(self):
        print("loaded")
        self.is_model_loaded = bool(True)
        if self.is_config_changed:
            self.is_config_changed = bool(False)
            self.is_model_loaded = bool(False)
            self.process_thread.start()
        elif self.is_queued:
            self.process_thread.start()

    def clear_opt_ticks(self):
        for opt in self.model_opts:
            opt.setIcon(self.noIcon)

    def method_opt11(self):
        self.clear_opt_ticks()
        self.opt11.setIcon(self.tick_icon)
        self.change_model(cnn_rnn_ctc_greedy_search)

    def method_opt12(self):
        self.clear_opt_ticks()
        self.opt12.setIcon(self.tick_icon)
        self.change_model(cnn_rnn_ctc_beam_search)

    def method_opt13(self):
        self.clear_opt_ticks()
        self.opt13.setIcon(self.tick_icon)
        self.change_model(cnn_rnn_ctc_beam_search_with_lm)

    def method_opt21(self):
        self.clear_opt_ticks()
        self.opt21.setIcon(self.tick_icon) 
        self.change_model(encoder_decoder_greedy_search)

    def method_opt22(self):
        self.clear_opt_ticks()
        self.opt22.setIcon(self.tick_icon) 
        self.change_model(encoder_decoder_beam_search)

    def change_config(self):
        
        fileName = QFileDialog.getOpenFileName(
            self,
            "Select configuration file",
            self.selected_folder,
            ("JSON Files (*.json)"),
        )[0]
        if fileName:
            self.clear_opt_ticks()
            self.cust_conf.setIcon(self.tick_icon)
            self.selected_folder = os.path.dirname(fileName)
            self.change_model(fileName)

    def change_model(self, name):
        self.statusBar().showMessage("Changing config file")
        self.process_thread.config_path = name
        self.process_thread.has_config_changed = bool(True)
        self.is_config_changed = bool(True)
        self.is_model_loaded = bool(False)
        self.process_thread.start()

    def sync_split(self):
        self.statusBar().showMessage("Done")
        if self.isSplit:
            self.ui.textBrowser.setPlainText(self.result_pages_list_g[self.currpic])
            self.changeicon(self.ui.split_pages_button, "&#64087", "black")
            self.ui.split_pages_button.setToolTip("Combine text")
            self.single_image_show()

        else:
            result_combined = "".join(self.result_pages_list_g)
            self.ui.textBrowser.setPlainText(result_combined)
            self.changeicon(self.ui.split_pages_button, "&#62970", "black")
            self.ui.split_pages_button.setToolTip("Split text to individual pages")

            self.single_image_hide()

    def scanning_complete(self, scanned_list):
        pyinsane2.exit()
        self.statusBar().showMessage("")
        self.scanAction.setEnabled(True)
        if scanned_list:
            print("scanning successful")
            self.newImageList(scanned_list)
        else:
            QtWidgets.QMessageBox.critical(
                self,
                "Scanner Not Found",
                "No scanner was detected.\nYour scanner may not be connected properly.\nOr your scanner model may not be supported (see www.sane-project.org/lists/sane-mfgs-cvs.html for a list of supported devices).",
                QtWidgets.QMessageBox.Ok,
            )

class AboutWindow(QtWidgets.QDialog):
    def __init__(self):
        super(AboutWindow, self).__init__()
        about_ui = abtform.Ui_Dialog()
        about_ui.setupUi(self)

        file = QFile(":/text/about.txt")
        file.open(QFile.ReadOnly | QFile.Text)
        stream = QTextStream(file)
        about_ui.text.setHtml(stream.readAll())


class TermsWindow(QtWidgets.QDialog):
    def __init__(self):
        super(TermsWindow, self).__init__()
        terms_ui = termsform.Ui_Dialog()
        terms_ui.setupUi(self)

        file = QFile(":/text/terms.txt")
        file.open(QFile.ReadOnly | QFile.Text)
        stream = QTextStream(file)
        terms_ui.textEdit.setHtml(stream.readAll())


class DocsWindow(QtWidgets.QDialog):
    def __init__(self):
        super(DocsWindow, self).__init__()
        docs_ui = docsform.Ui_Dialog()
        docs_ui.setupUi(self)

        file = QFile(":/text/docum.txt")
        file.open(QFile.ReadOnly | QFile.Text)
        stream = QTextStream(file)
        docs_ui.textEdit.setHtml(stream.readAll())

class ScanningThread(QtCore.QThread):
    scan_done = Signal(list)

    def __init__(self):
        QtCore.QThread.__init__(self)
        self.tempdir = ""

    def run(self):
        temp_img_name = "1.png"
        temp_img_path = os.path.join(self.tempdir, temp_img_name)
        new_list = []
        try:
            devices = pyinsane2.get_devices()
            if len(devices) <= 0:
                print("Scanning failed")
                self.scan_done.emit(None)
                return
            device = devices[0]
            print("Using scanner: %s" % (str(device)))

            pyinsane2.set_scanner_opt(device, "resolution", [300])
            pyinsane2.set_scanner_opt(device, "mode", ["Color"])

            pyinsane2.maximize_scan_area(device)
            # self.statusBar().showMessage("Scanning")
            scan_session = device.scan(multiple=False)
            try:
                while True:
                    scan_session.scan.read()
            except EOFError:
                pass
            image = scan_session.images[-1]
            image.save(temp_img_path, "PNG")
            new_list.append(temp_img_path)
            self.scan_done.emit(new_list)
        finally:
            print("exiting")


class ImagesScrollView(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super(ImagesScrollView, self).__init__(parent)
        self.vertlayout = QtWidgets.QVBoxLayout()
        self.setLayout(self.vertlayout)

    def dispImages(self, imageslist, w, h):
        self.clearLayout(self.vertlayout)
        for imgpath in imageslist:
            label = QtWidgets.QLabel(self)
            label.setStyleSheet(
                "border: 0px solid silver; border-bottom: 1px solid silver;"
            )
            self.vertlayout.addWidget(label)
            pixmap = QtGui.QPixmap(imgpath)
            label.setPixmap(pixmap.scaledToWidth(w - 35))

        vertspacer = QtWidgets.QSpacerItem(
            20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding
        )
        self.vertlayout.addSpacerItem(vertspacer)

    def clearLayout(self, layout):
        if layout != None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget() is not None:
                    child.widget().deleteLater()
                elif child.layout() is not None:
                    self.clearLayout(child.layout())

class ProcessingThread(QtCore.QThread):
    completed_signal = Signal(list)
    page_done = Signal(str)

    halted_signal = Signal()
    model_loaded_signal = Signal()

    def __init__(self):
        QtCore.QThread.__init__(self)
        self.deciphered = ""
        self.image_list = []
        self._isRunning = True
        self.model = None
        self.first_run = bool(True)
        self.has_config_changed = bool(False)
        self.config_path = default_config_file

    def run(self):
        if self.first_run:
            self.model, _ = create_and_run_model(self.config_path, eval_path=None, infer=True)
            self.model_loaded_signal.emit()
            self.first_run = bool(False)
            return

        if self.has_config_changed:
            self.has_config_changed = bool(False)
            print(self.config_path)
            self.model, _ = create_and_run_model(self.config_path, eval_path=None, infer=True)
            self.model_loaded_signal.emit()
            return

        if not self._isRunning:
            self._isRunning = True

        self.ocr_list = []
        for current_image_path in self.image_list:
            if self._isRunning == False:
                break

            _, model_output = self.model(current_image_path)
            pre_output = ""
            for mo in model_output:
                pre_output += mo
            self.ocr_list.append(pre_output)

        if self._isRunning == True:
            self.completed_signal.emit(self.ocr_list)
        else:
            self.halted_signal.emit()

    def halt(self):
        self._isRunning = False

def main():
    QtWidgets.QApplication.setStyle("fusion")
    app = QtWidgets.QApplication(sys.argv)
    mainWin = MainWindow()
    mainWin.show()
    sys.exit(app.exec_())
